#!/usr/bin/env python3
"""
Text extraction for PDF and Word files downloaded from Lark Drive.

Used by lark_sync's drive walker and by the submissions skill's
on-demand candidate-folder pull. Pure stdlib + pypdf + python-docx —
no external services. Dispatch by filename extension.

CLI: python tools/file_text.py <local-path>
"""

import io
import os
import sys
from typing import Optional


def _extract_pdf_text(blob: bytes) -> str:
    """PDF -> plain text via pypdf. Returns '' on failure (logged)."""
    try:
        from pypdf import PdfReader
    except ImportError:
        print("[file_text] pypdf not installed", file=sys.stderr)
        return ""
    try:
        reader = PdfReader(io.BytesIO(blob))
    except Exception as e:
        print(f"[file_text] pdf open failed: {e}", file=sys.stderr)
        return ""
    parts = []
    for p in reader.pages:
        try:
            t = p.extract_text() or ""
        except Exception:
            t = ""
        if t.strip():
            parts.append(t)
    return "\n\n".join(parts).strip()


def _extract_docx_text(blob: bytes) -> str:
    """Word (.docx) -> plain text via python-docx. Tables included."""
    try:
        from docx import Document
    except ImportError:
        print("[file_text] python-docx not installed", file=sys.stderr)
        return ""
    try:
        doc = Document(io.BytesIO(blob))
    except Exception as e:
        print(f"[file_text] docx open failed: {e}", file=sys.stderr)
        return ""
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract_file_text(blob: bytes, filename: str) -> Optional[str]:
    """Dispatch by extension. Returns None for unsupported (so caller
    can decide whether to skip silently or log)."""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _extract_pdf_text(blob)
    if name.endswith(".docx"):
        return _extract_docx_text(blob)
    if name.endswith(".doc"):
        # Legacy Word format — pure-python extractors don't work
        # reliably. Skip with a marker so the index knows.
        return ("(legacy .doc format — please re-save as .docx or PDF "
                "for ingestion)")
    return None


def main() -> int:
    if len(sys.argv) < 2:
        print(__doc__); return 0
    path = sys.argv[1]
    with open(path, "rb") as f:
        blob = f.read()
    text = extract_file_text(blob, path)
    if text is None:
        print(f"unsupported: {os.path.basename(path)}", file=sys.stderr)
        return 1
    print(text)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
